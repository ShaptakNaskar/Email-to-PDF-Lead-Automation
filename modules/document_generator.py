# ============================================================================
# MODULE: DOCUMENT GENERATOR
# ============================================================================
# Handles Word document creation and PDF conversion
# ============================================================================

import os
import re
import subprocess
from datetime import datetime
from docx import Document
from config import TEMPLATE_FILE, PERSONALISED_DIR, MASTER_LOG

def log_message(message: str, log_file: str = MASTER_LOG):
    """Log to console and file with timestamp."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = f"[{timestamp}] {message}\n"
    print(message)
    with open(log_file, 'a', encoding='utf-8') as logf:
        logf.write(log_entry)

def sanitize_filename(name):
    """Make filenames safe for all OS."""
    return re.sub(r'[\\/\*?:"<>|]', "_", name)

def replace_single_placeholder(doc, placeholder, replacement):
    """Replace a single placeholder in document paragraphs and runs."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

def replace_blurbs(doc, blurbs):
    """Replace blurb placeholders in document."""
    blurb_mappings = [
        ("Input Blerbs here", blurbs[0] if len(blurbs) > 0 else "Service offering"),
        ("Input Blerbs here", blurbs[1] if len(blurbs) > 1 else "Service offering"),
        ("Input Blerbs here", blurbs[2] if len(blurbs) > 2 else "Service offering"),
        ("Input Blerbs here", blurbs[3] if len(blurbs) > 3 else "Service offering"),
        ("Input Blerbs here", blurbs[4] if len(blurbs) > 4 else "Service offering"),
    ]

    # Replace in paragraphs
    found_count = 0
    for paragraph in doc.paragraphs:
        if found_count >= len(blurbs):
            break
        if "Input Blerbs here" in paragraph.text:
            for run in paragraph.runs:
                if "Input Blerbs here" in run.text:
                    run.text = run.text.replace("Input Blerbs here", blurbs[found_count])
                    found_count += 1
                    if found_count >= len(blurbs):
                        break

def generate_personalized_document(
    name: str,
    company_name: str,
    description: str,
    blurbs: list,
    message_id: str
) -> str:
    """
    Generate personalized DOCX document.
    Returns filename if successful, None if failed.
    """
    log_message(f"📄 Generating personalized document for {company_name}...")

    os.makedirs(PERSONALISED_DIR, exist_ok=True)

    if not os.path.exists(TEMPLATE_FILE):
        log_message(f"❌ Template '{TEMPLATE_FILE}' not found!")
        return None

    try:
        doc = Document(TEMPLATE_FILE)

        # Replace placeholders
        replace_single_placeholder(doc, "(Name)", name)
        replace_single_placeholder(doc, "(company name)", company_name)
        replace_single_placeholder(doc, "(what your company deals with)", description)
        replace_blurbs(doc, blurbs)

        # Generate filename
        safe_company = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')
        docx_filename = f"{safe_company}_{message_id[:8]}.docx"
        docx_path = os.path.join(PERSONALISED_DIR, docx_filename)

        doc.save(docx_path)
        log_message(f"✅ DOCX created: {docx_filename}")

        return docx_filename

    except Exception as e:
        log_message(f"❌ Error generating document: {type(e).__name__}: {str(e)}")
        return None

def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    Convert DOCX to PDF using LibreOffice.
    Returns PDF filename if successful, None if failed.
    """
    if not os.path.exists(docx_path):
        log_message(f"❌ DOCX file not found: {docx_path}")
        return None

    try:
        log_message(f"🔄 Converting DOCX to PDF...")

        # Try using LibreOffice
        command = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ]

        result = subprocess.run(command, capture_output=True, timeout=30)

        if result.returncode != 0:
            log_message(f"⚠️ LibreOffice conversion warning (code {result.returncode})")

        # Check if PDF was created
        pdf_filename = os.path.basename(docx_path).replace('.docx', '.pdf')
        pdf_path = os.path.join(output_dir, pdf_filename)

        if os.path.exists(pdf_path):
            log_message(f"✅ PDF created: {pdf_filename}")
            return pdf_filename
        else:
            log_message(f"❌ PDF conversion failed - file not found at {pdf_path}")
            return None

    except FileNotFoundError:
        log_message("❌ LibreOffice not found. Please install it:")
        log_message("   Windows: choco install libreoffice")
        log_message("   macOS: brew install libreoffice")
        log_message("   Linux: sudo apt-get install libreoffice")
        return None
    
    except subprocess.TimeoutExpired:
        log_message("❌ PDF conversion timed out (30 seconds)")
        return None
    
    except Exception as e:
        log_message(f"❌ Error converting to PDF: {type(e).__name__}: {str(e)}")
        return None