import os
import shutil
import subprocess
import time
import sys
from datetime import datetime
from docx import Document
from typing import List
import re

MASTER_LOG = "master_log.txt"
PERSONALISED_DIR = "personalised"
TEMPLATE_FILE = "template.docx"

def log_message(message: str, log_file: str = MASTER_LOG):
    """Log to both console and file with timestamp."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = f"[{timestamp}] {message}\n"
    print(message)
    with open(log_file, 'a', encoding='utf-8') as logf:
        logf.write(log_entry)

def replace_single_placeholder(doc: Document, placeholder: str, replacement: str):
    """Replace all occurrences of a single placeholder."""
    replaced = False
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            parts = paragraph.text.split(placeholder, 1)
            if len(parts) == 2:
                paragraph.text = parts[0] + replacement + parts[1]
                replaced = True
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        parts = paragraph.text.split(placeholder, 1)
                        if len(parts) == 2:
                            paragraph.text = parts[0] + replacement + parts[1]
                            replaced = True
    
    return replaced

def replace_blurbs(doc: Document, blurbs: List[str]):
    """Replace each 'Input Blurbs here' occurrence with next blurb."""
    blurb_idx = 0
    
    for paragraph in doc.paragraphs:
        if "Input Blurbs here" in paragraph.text and blurb_idx < len(blurbs):
            parts = paragraph.text.split("Input Blurbs here", 1)
            if len(parts) == 2:
                paragraph.text = parts[0] + blurbs[blurb_idx] + parts[1]
                blurb_idx += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Input Blurbs here" in paragraph.text and blurb_idx < len(blurbs):
                        parts = paragraph.text.split("Input Blurbs here", 1)
                        if len(parts) == 2:
                            paragraph.text = parts[0] + blurbs[blurb_idx] + parts[1]
                            blurb_idx += 1

def convert_docx_to_pdf(docx_path: str, pdf_dir: str, company_name: str) -> str:
    """
    Convert DOCX to PDF using LibreOffice or docx2pdf.
    """
    pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    pdf_path = os.path.join(pdf_dir, pdf_filename)
    
    log_message(f"🧾 Starting PDF conversion for: {docx_path}")
    
    # Check for LibreOffice
    soffice_cmd = shutil.which("soffice.exe") or shutil.which("soffice") or shutil.which("libreoffice")
    
    if soffice_cmd:
        log_message(f"📌 LibreOffice found at: {soffice_cmd}")
        cmd = [
            soffice_cmd, '--headless', '--convert-to', 'pdf',
            '--outdir', pdf_dir, docx_path
        ]
        
        try:
            process = subprocess.Popen(
                cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
            )
            
            start_time = time.time()
            while True:
                if process.poll() is not None:
                    break
                
                if os.path.exists(pdf_path):
                    log_message(f"✅ PDF detected at {round(time.time() - start_time, 1)}s")
                    break
                
                if time.time() - start_time > 60:
                    break
                
                time.sleep(1)
            
            try:
                process.communicate(timeout=5)
            except subprocess.TimeoutExpired:
                pass
            
            wait_time = 0
            while wait_time < 15 and not os.path.exists(pdf_path):
                time.sleep(1)
                wait_time += 1
            
            if os.path.exists(pdf_path):
                log_message(f"✅ Successfully created PDF via LibreOffice: {pdf_filename}")
                return pdf_filename
        
        except Exception as e:
            log_message(f"⚠️ LibreOffice conversion failed: {type(e).__name__}: {str(e)}")
    
    else:
        log_message("📌 LibreOffice not found. Trying docx2pdf as fallback...")
    
    # Fallback: Try docx2pdf
    try:
        import docx2pdf
    except ImportError:
        log_message("❌ Neither LibreOffice nor docx2pdf is installed!")
        
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        log_message(f"🧹 Deleted created DOCX: {docx_path}")
        print("\n❌ FATAL: Neither LibreOffice nor docx2pdf is available.")
        print("Please install one of the following:")
        print(" 1. LibreOffice (recommended) - add to PATH")
        print(" 2. docx2pdf (requires MS Word to be installed)")
        sys.exit(1)
    
    try:
        docx2pdf.convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            log_message(f"✅ Successfully created PDF via docx2pdf: {pdf_filename}")
            return pdf_filename
        else:
            return ""
    
    except Exception as e:
        error_str = str(e).lower()
        
        if any(keyword in error_str for keyword in ['word', 'not installed', 'could not locate', 'microsoft']):
            log_message("❌ docx2pdf found, but MS Word is NOT installed!")
            
            if os.path.exists(docx_path):
                os.remove(docx_path)
            
            log_message(f"🧹 Deleted created DOCX: {docx_path}")
            print("\n❌ FATAL: MS Word is not installed.")
            print("docx2pdf requires MS Word to work. Please install Microsoft Word and try again.")
            print("Script will exit now. No CSV updates or further actions will be performed.")
            sys.exit(1)
        
        else:
            log_message(f"❌ docx2pdf conversion failed: {type(e).__name__}: {str(e)}")
        
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        return ""

def generate_pdf(name: str, email: str, company_name: str, description: str, blurbs: List[str], message_id: str, sender_email: str) -> str:
    """
    Generate personalized PDF for the prospect.
    Returns filename if successful, empty string otherwise.
    """
    log_message(f"📄 Generating personalized PDF...")
    
    os.makedirs(PERSONALISED_DIR, exist_ok=True)
    
    if not os.path.exists(TEMPLATE_FILE):
        log_message(f"❌ Template '{TEMPLATE_FILE}' not found!")
        return ""
    
    try:
        doc = Document(TEMPLATE_FILE)
        
        replace_single_placeholder(doc, "(Name)", name)
        replace_single_placeholder(doc, "(company name)", company_name)
        replace_single_placeholder(doc, "(what your company deals with)", description)
        replace_blurbs(doc, blurbs)
        
        safe_company = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')
        docx_filename = f"{safe_company}_{message_id[:8]}.docx"
        docx_path = os.path.join(PERSONALISED_DIR, docx_filename)
        
        doc.save(docx_path)
        
        log_message(f"✅ DOCX created: {docx_filename}")
        
        pdf_filename = convert_docx_to_pdf(docx_path, PERSONALISED_DIR, company_name)
        
        if not pdf_filename:
            log_message(f"❌ PDF conversion failed.")
            return ""
        
        log_message(f"✅ PDF created: {pdf_filename}")
        return pdf_filename
    
    except Exception as e:
        log_message(f"❌ Error generating PDF: {type(e).__name__}: {str(e)}")
        return ""
